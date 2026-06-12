"""
Converte implantacao_shopfacil.md -> DOCX
Estilo: azul discreto nos titulos, tabelas formatadas
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SOURCE = r"R:\Arquivos\Codigos\PROJETO_TCC\Claude_Forencia\implantacao_shopfacil.md"
OUTPUT = r"R:\Arquivos\Trabalho Rubens\implantacao_shopfacil.docx"

# Cores (RGB)
AZUL_TITULO  = RGBColor(44,  82,  130)
AZUL_H3      = RGBColor(70,  110, 160)
AZUL_HDR_TAB = RGBColor(44,  82,  130)
AZUL_PAR_TAB = RGBColor(219, 232, 247)
PRETO        = RGBColor(0,   0,   0)
BRANCO       = RGBColor(255, 255, 255)
CINZA_ROD    = RGBColor(120, 120, 120)

CHAR_MAP = {
    "—": "--", "–": "-",
    "‘": "'",  "’": "'",
    "“": '"',  "”": '"',
    " ": " ",
}

def sanitize(text):
    for ch, rep in CHAR_MAP.items():
        text = text.replace(ch, rep)
    return text


# ── helpers XML para cor de fundo em célula ──────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def rgb_to_hex(rgb: RGBColor):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_border(cell, color_hex="A0B4DC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tcBorders.append(el)
    tcPr.append(tcBorders)


# ── setup estilos do documento ───────────────────────────────────────────────
def setup_doc():
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Estilo Normal base
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = PRETO

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name  = "Arial"
    h1.font.size  = Pt(14)
    h1.font.bold  = True
    h1.font.color.rgb = BRANCO
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after  = Pt(6)
    # Fundo azul via shading é feito parágrafo a parágrafo

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name  = "Arial"
    h2.font.size  = Pt(12)
    h2.font.bold  = True
    h2.font.color.rgb = AZUL_TITULO
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after  = Pt(4)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name  = "Arial"
    h3.font.size  = Pt(10.5)
    h3.font.bold  = True
    h3.font.color.rgb = AZUL_H3
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after  = Pt(2)

    # Heading 4
    h4 = doc.styles["Heading 4"]
    h4.font.name   = "Arial"
    h4.font.size   = Pt(10)
    h4.font.bold   = True
    h4.font.italic = True
    h4.font.color.rgb = AZUL_H3

    return doc


def shade_paragraph(para, hex_color):
    """Aplica fundo colorido a um parágrafo (usado no H1)."""
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def add_paragraph_border_bottom(para, color_hex="96B4DC"):
    """Adiciona borda inferior no parágrafo (sublinhado H2)."""
    pPr    = para._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── renderiza inline bold / code ─────────────────────────────────────────────
def add_inline(para, text, base_size=10, base_color=PRETO):
    partes = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for p in partes:
        if p.startswith("**") and p.endswith("**"):
            run = para.add_run(p[2:-2])
            run.bold = True
            run.font.size  = Pt(base_size)
            run.font.color.rgb = base_color
        elif p.startswith("`") and p.endswith("`"):
            run = para.add_run(p[1:-1])
            run.font.name  = "Courier New"
            run.font.size  = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(150, 100, 0)
        else:
            run = para.add_run(p)
            run.font.size  = Pt(base_size)
            run.font.color.rgb = base_color


# ── tabela ───────────────────────────────────────────────────────────────────
def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines):
        ln = lines[i].strip()
        if not ln.startswith("|"):
            break
        if re.match(r'^\|[-|: ]+\|$', ln):
            i += 1
            continue
        cells = [c.strip() for c in ln.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)

    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    hdr_hex = rgb_to_hex(AZUL_HDR_TAB)
    par_hex = rgb_to_hex(AZUL_PAR_TAB)

    for ri, row in enumerate(rows):
        is_hdr = ri == 0
        bg_hex = hdr_hex if is_hdr else (par_hex if ri % 2 == 0 else "FFFFFF")
        for ci in range(ncols):
            cell = tbl.cell(ri, ci)
            text = row[ci] if ci < len(row) else ""
            set_cell_bg(cell, bg_hex)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_hdr else WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            run = para.add_run(text)
            run.font.name  = "Arial"
            run.font.size  = Pt(8.5)
            run.font.bold  = is_hdr
            run.font.color.rgb = BRANCO if is_hdr else PRETO

    doc.add_paragraph()  # espaco apos tabela


# ── capa ──────────────────────────────────────────────────────────────────────
def add_capa(doc):
    # Espacamento topo
    for _ in range(6):
        doc.add_paragraph()

    # Linha superior
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(0)
    pPr  = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "12")
    top.set(qn("w:color"), "2C5282")
    pBdr.append(top)
    pPr.append(pBdr)

    # Titulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("Projeto de Implantacao de Software ERP")
    run.font.name  = "Arial"
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = AZUL_TITULO

    # Empresa
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    run2 = p2.add_run("ShopFacil E-Commerce Ltda.")
    run2.font.name   = "Arial"
    run2.font.size   = Pt(14)
    run2.font.italic = True
    run2.font.color.rgb = AZUL_H3

    # Linha inferior
    p_line2 = doc.add_paragraph()
    p_line2.paragraph_format.space_after = Pt(16)
    pPr2  = p_line2._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    bot   = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12")
    bot.set(qn("w:color"), "2C5282")
    pBdr2.append(bot)
    pPr2.append(pBdr2)

    # Dados
    dados = [
        ("Software Escolhido", "Omie ERP -- SaaS Cloud"),
        ("Orcamento Total",    "R$ 50.000,00"),
        ("Duracao do Projeto", "12 meses implantacao + 12 meses suporte"),
        ("Equipe",             "5 integrantes"),
        ("Emissao",            "Maio de 2026"),
    ]
    for label, valor in dados:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(label + ": ")
        r1.font.name = "Arial"
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = AZUL_TITULO
        r2 = p.add_run(valor)
        r2.font.name = "Arial"
        r2.font.size = Pt(10)
        r2.font.color.rgb = PRETO

    # Quebra de pagina apos capa
    doc.add_page_break()


# ── main ─────────────────────────────────────────────────────────────────────
def main():
    with open(SOURCE, encoding="utf-8") as f:
        lines = [sanitize(ln.rstrip("\n")) for ln in f]

    doc = setup_doc()
    add_capa(doc)

    skip_header  = True
    skipped      = 0
    i            = 0

    while i < len(lines):
        line    = lines[i]
        stripped = line.strip()

        # Pula cabecalho MD (capa ja desenhada)
        if skip_header and skipped < 10:
            skipped += 1
            i += 1
            continue
        skip_header = False

        # Separador ---
        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            add_paragraph_border_bottom(p, "CCCCCC")
            i += 1
            continue

        # H1
        if re.match(r'^# ', stripped):
            texto = stripped[2:]
            p = doc.add_heading(texto, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            shade_paragraph(p, "2C5282")
            for run in p.runs:
                run.font.color.rgb = BRANCO
            i += 1
            continue

        # H2
        if re.match(r'^## ', stripped):
            texto = stripped[3:]
            p = doc.add_heading(texto, level=2)
            add_paragraph_border_bottom(p, "96B4DC")
            i += 1
            continue

        # H3
        if re.match(r'^### ', stripped):
            texto = stripped[4:]
            doc.add_heading(texto, level=3)
            i += 1
            continue

        # H4
        if re.match(r'^#### ', stripped):
            texto = stripped[5:]
            doc.add_heading(texto, level=4)
            i += 1
            continue

        # Tabela
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        # Bloco de codigo
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            for cl in code_lines:
                p = doc.add_paragraph(cl)
                p.paragraph_format.left_indent = Cm(0.8)
                for run in p.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(60, 60, 60)
            continue

        # Bullet
        if re.match(r'^[-*] ', stripped):
            texto = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, texto)
            i += 1
            continue

        # Lista numerada
        if re.match(r'^\d+\. ', stripped):
            texto = re.sub(r'^\d+\. ', '', stripped)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, texto)
            i += 1
            continue

        # Linha vazia
        if not stripped:
            i += 1
            continue

        # Paragrafo normal
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline(p, stripped)
        i += 1

    doc.save(OUTPUT)
    print("DOCX gerado em: " + OUTPUT)


if __name__ == "__main__":
    main()
